"""Generate Results & Analysis chapter and IEEE Research Paper as .docx files with embedded charts."""
import re, os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ARTIFACTS = r"C:\Users\sadaf\.gemini\antigravity\brain\5a34dd0f-9433-4bd6-9fa2-b25cb52d43a0\artifacts"
CHARTS = r"C:\Users\sadaf\OneDrive\Desktop\ornamis-2\doc_assets\results"
OUT1 = r"C:\Users\sadaf\OneDrive\Desktop\ORNAMIS_Results_Analysis.docx"
OUT2 = r"C:\Users\sadaf\OneDrive\Desktop\ORNAMIS_Research_Paper_IEEE.docx"

CHART_MAP = {
    'chart_ai_response_time.png': 'chart_ai_response_time.png',
    'chart_ai_success_rate.png': 'chart_ai_success_rate.png',
    'chart_ar_fps_browsers.png': 'chart_ar_fps_browsers.png',
    'chart_ar_detection_accuracy.png': 'chart_ar_detection_accuracy.png',
    'chart_gold_api_response.png': 'chart_gold_api_response.png',
    'chart_gold_price_accuracy.png': 'chart_gold_price_accuracy.png',
    'chart_lighthouse_scores.png': 'chart_lighthouse_scores.png',
    'chart_page_load_times.png': 'chart_page_load_times.png',
    'chart_user_survey_radar.png': 'chart_user_survey_radar.png',
    'chart_test_coverage.png': 'chart_test_coverage.png',
    'chart_feature_comparison.png': 'chart_feature_comparison.png',
    'chart_fallback_activation.png': 'chart_fallback_activation.png',
}

FIGURE_TO_CHART = {
    'Figure R.1': 'chart_ai_response_time.png',
    'Figure R.2': 'chart_ai_success_rate.png',
    'Figure R.3': 'chart_ar_fps_browsers.png',
    'Figure R.4': 'chart_ar_detection_accuracy.png',
    'Figure R.5': 'chart_gold_api_response.png',
    'Figure R.6': 'chart_gold_price_accuracy.png',
    'Figure R.7': 'chart_lighthouse_scores.png',
    'Figure R.8': 'chart_page_load_times.png',
    'Figure R.9': 'chart_user_survey_radar.png',
    'Figure R.10': 'chart_test_coverage.png',
    'Figure R.11': 'chart_feature_comparison.png',
    'Figure R.12': 'chart_fallback_activation.png',
}

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def style_doc(doc):
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(2.54)
    for i in range(1, 4):
        h = doc.styles[f'Heading {i}']
        h.font.name = 'Times New Roman'
        h.font.color.rgb = RGBColor(0x49, 0x2f, 0x0e)
        h.font.bold = True
        if i == 1:
            h.font.size = Pt(18)
            h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif i == 2:
            h.font.size = Pt(14)
        else:
            h.font.size = Pt(12)

def add_table(doc, header, rows):
    cols = len(header)
    table = doc.add_table(rows=1 + len(rows), cols=cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, text in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(text.strip())
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '492f0e')
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for r, row_data in enumerate(rows):
        for c, text in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            run = cell.paragraphs[0].add_run(text.strip())
            run.font.size = Pt(9)
            run.font.name = 'Times New Roman'
    doc.add_paragraph()

def add_code(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def process_inline(para, text):
    parts = re.split(r'(\*\*.*?\*\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = para.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
        else:
            run = para.add_run(part)
        run.font.name = run.font.name or 'Times New Roman'
        run.font.size = run.font.size or Pt(12)

def try_insert_chart(doc, line):
    """Check if a blockquote line references a figure and insert the chart image."""
    for fig_key, chart_file in FIGURE_TO_CHART.items():
        if fig_key in line:
            chart_path = os.path.join(CHARTS, chart_file)
            if os.path.exists(chart_path):
                # Add figure caption
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(line.lstrip('> ').strip())
                run.bold = True
                run.italic = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x49, 0x2f, 0x0e)
                # Insert image
                doc.add_picture(chart_path, width=Inches(5.5))
                last_p = doc.paragraphs[-1]
                last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()  # spacing
                return True
    return False

def parse_md(doc, text):
    lines = text.split('\n')
    i = 0
    in_code = False
    code_buf = []
    while i < len(lines):
        line = lines[i]
        if line.strip().startswith('```'):
            if in_code:
                add_code(doc, '\n'.join(code_buf))
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        if stripped == '---':
            doc.add_page_break()
            i += 1
            continue
        if stripped.startswith('# ') and not stripped.startswith('## '):
            doc.add_heading(stripped[2:].strip(), level=1)
            i += 1
            continue
        if stripped.startswith('## '):
            doc.add_heading(stripped[3:].strip(), level=2)
            i += 1
            continue
        if stripped.startswith('### '):
            doc.add_heading(stripped[4:].strip(), level=3)
            i += 1
            continue
        # Tables
        if '|' in stripped and stripped.startswith('|'):
            rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                row_text = lines[i].strip()
                cells = [c.strip() for c in row_text.split('|')[1:-1]]
                if cells and all(re.match(r'^[-:]+$', c) for c in cells):
                    i += 1
                    continue
                rows.append(cells)
                i += 1
            if len(rows) >= 2:
                add_table(doc, rows[0], rows[1:])
            elif len(rows) == 1:
                add_table(doc, rows[0], [])
            continue
        # Blockquotes - try to insert chart images
        if stripped.startswith('>'):
            text_content = stripped.lstrip('> ').strip()
            # Skip "Chart:" and "Note:" lines
            if text_content.startswith('*Chart:*') or text_content.startswith('*Note:*'):
                i += 1
                continue
            if try_insert_chart(doc, stripped):
                i += 1
                continue
            # Regular blockquote
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            process_inline(p, text_content)
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x49, 0x2f, 0x0e)
            i += 1
            continue
        if re.match(r'^\d+\.', stripped):
            text_content = re.sub(r'^\d+\.\s*', '', stripped)
            p = doc.add_paragraph(style='List Number')
            process_inline(p, text_content)
            i += 1
            continue
        if stripped.startswith('- ') or stripped.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            process_inline(p, stripped[2:].strip())
            i += 1
            continue
        p = doc.add_paragraph()
        process_inline(p, stripped)
        i += 1

def generate_doc(md_path, out_path, title):
    doc = Document()
    style_doc(doc)
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    parse_md(doc, content)
    doc.save(out_path)
    print(f"  [OK] {title} -> {out_path}")

if __name__ == '__main__':
    print("=" * 55)
    print("  Generating DOCX files with embedded charts")
    print("=" * 55)
    
    results_md = os.path.join(ARTIFACTS, "results_analysis_chapter.md")
    paper_md = os.path.join(ARTIFACTS, "research_paper_ieee.md")
    
    generate_doc(results_md, OUT1, "Results & Analysis Chapter")
    generate_doc(paper_md, OUT2, "IEEE Research Paper")
    
    print(f"\n[DONE] Both files saved to Desktop!")
    print(f"  1. ORNAMIS_Results_Analysis.docx")
    print(f"  2. ORNAMIS_Research_Paper_IEEE.docx")
