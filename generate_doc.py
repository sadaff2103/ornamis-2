"""Generate ORNAMIS project report as a .docx file from markdown chapter files."""
import re
# pyrefly: ignore [missing-import]
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

ARTIFACTS_DIR = r"C:\Users\sadaf\.gemini\antigravity\brain\5a34dd0f-9433-4bd6-9fa2-b25cb52d43a0\artifacts"
OUTPUT_PATH = r"C:\Users\sadaf\OneDrive\Desktop\ORNAMIS_Report_10Chapters.docx"

FILES = [
    os.path.join(ARTIFACTS_DIR, "chapter_1_2_introduction.md"),
    os.path.join(ARTIFACTS_DIR, "chapter_3_system_analysis.md"),
    os.path.join(ARTIFACTS_DIR, "chapter_4_system_design.md"),
    os.path.join(ARTIFACTS_DIR, "chapter_5_implementation.md"),
    os.path.join(ARTIFACTS_DIR, "chapter_6_results.md"),
    os.path.join(ARTIFACTS_DIR, "chapter_7_screenshots.md"),
    os.path.join(ARTIFACTS_DIR, "chapter_8_testing.md"),
    os.path.join(ARTIFACTS_DIR, "chapter_9_future.md"),
    os.path.join(ARTIFACTS_DIR, "chapter_10_conclusion.md"),
]

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def style_doc(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    
    for i in range(1, 4):
        h = doc.styles[f'Heading {i}']
        h.font.name = 'Times New Roman'
        h.font.color.rgb = RGBColor(0x49, 0x2f, 0x0e)
        h.font.bold = True
        if i == 1:
            h.font.size = Pt(18)
            h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            h.paragraph_format.space_before = Pt(24)
            h.paragraph_format.space_after = Pt(12)
        elif i == 2:
            h.font.size = Pt(14)
            h.paragraph_format.space_before = Pt(18)
            h.paragraph_format.space_after = Pt(8)
        else:
            h.font.size = Pt(12)
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(6)

    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(2.54)

def add_table(doc, header_row, data_rows):
    cols = len(header_row)
    table = doc.add_table(rows=1 + len(data_rows), cols=cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header
    for i, text in enumerate(header_row):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(text.strip())
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '492f0e')
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    
    # Data rows
    for r, row_data in enumerate(data_rows):
        for c, text in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(text.strip())
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
    
    doc.add_paragraph()

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def process_inline(doc_paragraph, text):
    """Process inline markdown bold/code in a paragraph."""
    parts = re.split(r'(\*\*.*?\*\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = doc_paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        elif part.startswith('`') and part.endswith('`'):
            run = doc_paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
        else:
            run = doc_paragraph.add_run(part)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

def parse_md_to_docx(doc, md_text):
    lines = md_text.split('\n')
    i = 0
    in_code = False
    code_buf = []
    
    while i < len(lines):
        line = lines[i]
        
        # Code blocks
        if line.strip().startswith('```'):
            if in_code:
                add_code_block(doc, '\n'.join(code_buf))
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        
        if in_code:
            code_buf.append(line)
            i += 1
            continue
        
        stripped = line.strip()
        
        # Skip empty lines
        if not stripped:
            i += 1
            continue
        
        # Skip horizontal rules
        if stripped == '---':
            doc.add_page_break()
            i += 1
            continue
        
        # Headings
        if stripped.startswith('# ') and not stripped.startswith('## '):
            title = stripped[2:].strip()
            doc.add_heading(title, level=1)
            i += 1
            continue
        
        if stripped.startswith('## '):
            title = stripped[3:].strip()
            doc.add_heading(title, level=2)
            i += 1
            continue
        
        if stripped.startswith('### '):
            title = stripped[4:].strip()
            doc.add_heading(title, level=3)
            i += 1
            continue
        
        # Tables
        if '|' in stripped and stripped.startswith('|'):
            rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                row_text = lines[i].strip()
                cells = [c.strip() for c in row_text.split('|')[1:-1]]
                # Skip separator rows
                if cells and all(re.match(r'^[-:]+$', c) for c in cells):
                    i += 1
                    continue
                rows.append(cells)
                i += 1
            
            if len(rows) >= 2:
                add_table(doc, rows[0], rows[1:])
            elif len(rows) == 1:
                add_table(doc, rows[0], [])
            continue
        
        # Blockquotes (figure references)
        if stripped.startswith('>'):
            text = stripped.lstrip('> ').strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_before = Pt(6)
            process_inline(p, text)
            for run in p.runs:
                run.font.italic = True
                run.font.color.rgb = RGBColor(0x49, 0x2f, 0x0e)
            i += 1
            continue
        
        # Numbered lists
        if re.match(r'^\d+\.', stripped):
            text = re.sub(r'^\d+\.\s*', '', stripped)
            p = doc.add_paragraph(style='List Number')
            process_inline(p, text)
            i += 1
            continue
        
        # Bullet lists
        if stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            process_inline(p, text)
            i += 1
            continue
        
        # Regular paragraph
        p = doc.add_paragraph()
        process_inline(p, stripped)
        i += 1

def main():
    doc = Document()
    style_doc(doc)
    
    # Read and combine all markdown files
    for idx, filepath in enumerate(FILES):
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()
        
        if idx > 0:
            doc.add_page_break()
        
        parse_md_to_docx(doc, content)
    
    doc.save(OUTPUT_PATH)
    print(f"Document saved to: {OUTPUT_PATH}")

if __name__ == '__main__':
    main()
